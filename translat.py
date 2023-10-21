import re
from docx import Document

persian_path = 'Persian.docx'
path_1 = 'zirnevis.docx'

try:
    doc_2 = Document(path_1)
    doc_1 = Document(persian_path)
    
    # for line_persian in doc_1.paragraphs:
    #     print(line_persian.text)
    for line in doc_2.paragraphs:
        for line_text in line.text:
            if re.search('[a-zA-Z]', line_text):
                # line = line_persian
                print(line_text)
except Exception as e:
    print(f'error:{e}')




# from docx import Document

# # مسیر فایل Word
# file_path = 'path_to_your_word_file.docx'

# # تلاش برای باز کردن فایل
# try:
#     doc = Document(file_path)
#     for para in doc.paragraphs:
#         print(para.text)  # یا انجام عملیات دیگر با هر خط
# except Exception as e:
#     print(f"خطا در باز کردن فایل: {e}")

                 